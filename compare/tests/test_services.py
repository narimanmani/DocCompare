from __future__ import annotations

from io import BytesIO
import zipfile

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.shared import OxmlElement, qn

from compare.diff_utils import DiffComputationResult
from compare.diff_links import compare_links_by_text
from compare.services import DiffOptions, parse_docx_bytes, perform_comparison, store_diff_result
from compare.url_utils import UrlNormalizationOptions


def make_docx(paragraphs: list[str]) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def make_docx_with_link(url: str) -> bytes:
    document = Document()
    paragraph = document.add_paragraph("See")
    paragraph.add_run(" ")
    _add_hyperlink(paragraph, url, "Technology Test Strategy - FR")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_tracked_link(url: str) -> bytes:
    base = make_docx_with_link(url)
    source = BytesIO(base)
    buffer = BytesIO()
    with zipfile.ZipFile(source) as source_zip:
        with zipfile.ZipFile(buffer, "w") as target_zip:
            for info in source_zip.infolist():
                data = source_zip.read(info.filename)
                if info.filename == "word/document.xml":
                    xml = data.decode("utf-8")
                    xml = xml.replace("<w:hyperlink", "<w:ins><w:hyperlink", 1)
                    xml = xml.replace("</w:hyperlink>", "</w:hyperlink></w:ins>", 1)
                    data = xml.encode("utf-8")
                target_zip.writestr(info.filename, data)
    return buffer.getvalue()


@pytest.mark.django_db
def test_store_diff_result_creates_record() -> None:
    diff = DiffComputationResult(
        html="<p>ok</p>",
        summary={
            "insertions": 1,
            "deletions": 0,
            "replacements": 0,
            "percent_changed": 50.0,
            "link_text_matches": 0,
            "link_destination_changes": 0,
            "link_destination_matches": 0,
        },
        paragraphs=[],
        link_changes=[],
    )
    result = store_diff_result(diff, DiffOptions(), ["a.docx", "b.docx"])
    assert result.summary["insertions"] == 1
    assert result.options["ignore_case"] is True


@pytest.mark.django_db
def test_perform_comparison_generates_summary() -> None:
    file_a = make_docx(["Hello world", "Goodbye"])
    file_b = make_docx(["hello world", "Farewell"])

    diff, filenames = perform_comparison(
        [("first.docx", file_a), ("second.docx", file_b)],
        DiffOptions(ignore_case=True, ignore_punctuation=False, ignore_whitespace=False),
    )

    assert filenames == ["first.docx", "second.docx"]
    assert diff.summary["insertions"] >= 0
    assert diff.summary["percent_changed"] >= 0
    assert {"link_text_matches", "link_destination_changes", "link_destination_matches"} <= diff.summary.keys()


def test_docx_fallback_preserves_hyperlinks(monkeypatch) -> None:
    from compare import services

    monkeypatch.setattr(services, "_parse_with_pandoc", lambda path: None)

    first_url = "https://navcanada.sharepoint.com/sites/TCoE/Shared Documents/Testing Standards/Test Policy and Test Strategy/Technology Test Strategy - FR.pdf"
    second_url = "https://navcanada.sharepoint.com/sites/TCoE/Shared Documents/Testing Standards/Test Strategy/Technology Test Strategy - FR.pdf"

    left = parse_docx_bytes(make_docx_with_link(first_url))
    right = parse_docx_bytes(make_docx_with_link(second_url))

    assert left and right
    assert any(token.type == "anchor" for token in left[0].tokens)
    assert any(token.type == "anchor" for token in right[0].tokens)

    options = UrlNormalizationOptions(
        ignore_protocol=False,
        normalize_trailing_slash=False,
        lowercase_host=False,
        drop_tracking_params=False,
        strip_fragment=False,
    )

    left_statuses, right_statuses, records = compare_links_by_text(left, right, options)

    assert records
    assert records[0].text == "Technology Test Strategy - FR"
    assert records[0].left_href == first_url
    assert records[0].right_href == second_url
    assert records[0].changed is True
    assert any(status == "link-href-changed" for status in left_statuses.get(0, {}).values())
    assert any(status == "link-href-changed" for status in right_statuses.get(0, {}).values())


def test_python_docx_fallback_preserves_hyperlinks(monkeypatch) -> None:
    from compare import services

    monkeypatch.setattr(services, "_parse_with_pandoc", lambda path: None)
    monkeypatch.setattr(services, "_parse_docx_xml", lambda path: None)

    url = "https://example.com/docs"

    paragraphs = parse_docx_bytes(make_docx_with_link(url))

    assert paragraphs
    anchors = [token for token in paragraphs[0].tokens if token.type == "anchor"]
    assert anchors
    assert anchors[0].href == url


def test_parse_docx_accepts_tracked_changes(monkeypatch) -> None:
    from compare import services

    monkeypatch.setattr(services, "_parse_with_pandoc", lambda path: None)

    url = "https://example.com/docs"

    paragraphs = parse_docx_bytes(make_docx_with_tracked_link(url))

    assert paragraphs
    anchors = [token for token in paragraphs[0].tokens if token.type == "anchor"]
    assert anchors
    assert anchors[0].href == url
