"""Service layer for document comparison and persistence."""
from __future__ import annotations

import json
import logging
import subprocess
import zipfile
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence
from xml.etree import ElementTree as ET

from django.db import transaction
from django.utils.html import escape

from .diff_utils import DiffComputationResult, DiffOptions as EngineOptions, build_diff
from .html_tokens import Paragraph, Token, paragraphs_from_html, paragraphs_from_text
from .models import DiffResult
from .url_utils import UrlNormalizationOptions


logger = logging.getLogger(__name__)


@dataclass(slots=True)
class DiffOptions:
    """Options toggled by the user for diff normalization."""

    ignore_case: bool = True
    ignore_punctuation: bool = False
    ignore_whitespace: bool = False
    ignore_protocol: bool = False
    normalize_trailing_slash: bool = False
    drop_tracking_params: bool = True
    lowercase_host: bool = False
    ignore_url_fragments: bool = False

    @classmethod
    def from_dict(cls, data: dict[str, bool]) -> "DiffOptions":
        """Instantiate from a dictionary of request data."""
        return cls(
            ignore_case=bool(data.get("ignore_case", False)),
            ignore_punctuation=bool(data.get("ignore_punctuation", False)),
            ignore_whitespace=bool(data.get("ignore_whitespace", False)),
            ignore_protocol=bool(data.get("ignore_protocol", False)),
            normalize_trailing_slash=bool(data.get("normalize_trailing_slash", False)),
            drop_tracking_params=bool(data.get("drop_tracking_params", False)),
            lowercase_host=bool(data.get("lowercase_host", False)),
            ignore_url_fragments=bool(data.get("ignore_url_fragments", False)),
        )

    def as_dict(self) -> dict[str, bool]:
        return {
            "ignore_case": self.ignore_case,
            "ignore_punctuation": self.ignore_punctuation,
            "ignore_whitespace": self.ignore_whitespace,
            "ignore_protocol": self.ignore_protocol,
            "normalize_trailing_slash": self.normalize_trailing_slash,
            "drop_tracking_params": self.drop_tracking_params,
            "lowercase_host": self.lowercase_host,
            "ignore_url_fragments": self.ignore_url_fragments,
        }

    def url_options(self) -> UrlNormalizationOptions:
        return UrlNormalizationOptions(
            ignore_protocol=self.ignore_protocol,
            normalize_trailing_slash=self.normalize_trailing_slash,
            lowercase_host=self.lowercase_host,
            drop_tracking_params=self.drop_tracking_params,
            strip_fragment=self.ignore_url_fragments,
        )


class DocumentParseError(RuntimeError):
    """Raised when a document cannot be parsed."""


def _parse_with_docx2python(path: Path) -> list[str]:
    from docx2python import docx2python

    paragraphs: list[str] = []
    with docx2python(path) as doc:
        for block in doc.body:
            paragraphs.extend(_flatten_block(block))
    if getattr(doc, "images", None):  # type: ignore[attr-defined]
        for image in doc.images:  # type: ignore[attr-defined]
            paragraphs.append(f"[image] {Path(image).name}")
    return paragraphs


def _flatten_block(block: object) -> list[str]:
    if isinstance(block, str):
        text = block.strip()
        return [text] if text else []
    if isinstance(block, (list, tuple)):
        if block and any(isinstance(item, (list, tuple)) for item in block):
            rows: list[str] = []
            for item in block:
                cell = " ".join(_flatten_block(item))
                if cell:
                    rows.append(cell)
            if rows:
                return [f"[table] {' / '.join(rows)}"]
            return []
        texts: list[str] = []
        for item in block:
            texts.extend(_flatten_block(item))
        return texts
    return []


def _parse_with_python_docx(path: Path) -> list[str]:
    from docx import Document

    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        cells = []
        for row in table.rows:
            cells.append(" | ".join(cell.text.strip() for cell in row.cells))
        if cells:
            paragraphs.append(f"[table] {' / '.join(filter(None, cells))}")
    return paragraphs


def _parse_with_pandoc(path: Path) -> list[Paragraph] | None:
    try:
        completed = subprocess.run(
            ["pandoc", str(path), "-t", "html"],
            check=True,
            capture_output=True,
            text=True,
        )
    except FileNotFoundError:
        logger.warning("Pandoc is not installed; falling back to plain-text parsing.")
        return None
    except subprocess.CalledProcessError as exc:
        logger.warning("Pandoc conversion failed: %s", exc.stderr or exc)
        return None

    paragraphs = paragraphs_from_html(completed.stdout)
    if not paragraphs:
        logger.warning("Pandoc produced no paragraphs; falling back to plain-text parsing.")
        return None
    return paragraphs


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_TEXT_TAG = f"{{{_W_NS}}}t"
_TAB_TAG = f"{{{_W_NS}}}tab"
_BR_TAGS = {f"{{{_W_NS}}}br", f"{{{_W_NS}}}cr"}
_PARAGRAPH_PROPERTIES_TAG = f"{{{_W_NS}}}pPr"
_HYPERLINK_TAG = f"{{{_W_NS}}}hyperlink"
_ANCHOR_ATTR = f"{{{_W_NS}}}anchor"
_ID_ATTR = f"{{{_R_NS}}}id"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _normalize_docx_text(value: str) -> str:
    return " ".join(value.split())


def _extract_text(element: ET.Element) -> str:
    parts: list[str] = []
    for node in element.iter():
        if node.tag == _TEXT_TAG and node.text:
            parts.append(node.text)
        elif node.tag == _TAB_TAG:
            parts.append(" ")
        elif node.tag in _BR_TAGS:
            parts.append(" ")
    return _normalize_docx_text(" ".join(parts))


def _build_anchor_token(element: ET.Element, relationships: dict[str, str]) -> Token | None:
    rel_id = element.get(_ID_ATTR, "")
    href = relationships.get(rel_id, "") if rel_id else ""
    anchor = element.get(_ANCHOR_ATTR, "")
    if anchor and not href:
        href = f"#{anchor}"
    text = _extract_text(element)
    if not text:
        return None
    return Token(type="anchor", text=text, href=href)


def _parse_docx_xml(path: Path) -> list[Paragraph] | None:
    try:
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read("word/document.xml")
            rels_xml = archive.read("word/_rels/document.xml.rels")
    except KeyError:
        return None
    except FileNotFoundError:
        return None

    try:
        document_tree = ET.fromstring(document_xml)
        rels_tree = ET.fromstring(rels_xml)
    except ET.ParseError:
        return None

    relationships: dict[str, str] = {}
    for rel in rels_tree.findall(f"{{{_REL_NS}}}Relationship"):
        rel_type = rel.get("Type", "")
        if rel_type.endswith("/hyperlink"):
            rel_id = rel.get("Id")
            target = rel.get("Target", "")
            if rel_id:
                relationships[rel_id] = target

    paragraphs: list[Paragraph] = []
    for paragraph in document_tree.findall(".//w:p", {"w": _W_NS}):
        tokens: list[Token] = []
        text_buffer: list[str] = []

        def flush_text_buffer() -> None:
            if not text_buffer:
                return
            text = _normalize_docx_text(" ".join(text_buffer))
            text_buffer.clear()
            if text:
                tokens.append(Token(type="text", text=text))

        for child in paragraph:
            if child.tag == _PARAGRAPH_PROPERTIES_TAG:
                continue
            if child.tag == _HYPERLINK_TAG:
                flush_text_buffer()
                anchor = _build_anchor_token(child, relationships)
                if anchor:
                    tokens.append(anchor)
                continue

            text = _extract_text(child)
            if text:
                text_buffer.append(text)

        flush_text_buffer()

        meaningful = False
        if tokens:
            meaningful = any(token.type == "anchor" or token.text.strip() for token in tokens)
        if meaningful:
            paragraphs.append(Paragraph(tokens=tokens))

    if not paragraphs:
        return None
    return paragraphs


def parse_docx_bytes(data: bytes) -> list[Paragraph]:
    """Parse a DOCX file into structured paragraphs."""

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        tmp_file.write(data)
        tmp_path = Path(tmp_file.name)
    try:
        paragraphs = _parse_with_pandoc(tmp_path)
        if paragraphs is not None:
            return paragraphs

        xml_paragraphs = _parse_docx_xml(tmp_path)
        if xml_paragraphs is not None:
            return xml_paragraphs

        try:
            text_blocks = _parse_with_docx2python(tmp_path)
        except Exception:
            text_blocks = _parse_with_python_docx(tmp_path)
        return paragraphs_from_text(text_blocks)
    finally:
        tmp_path.unlink(missing_ok=True)


def parse_multiple(files: Iterable[tuple[str, bytes]]) -> tuple[list[Paragraph], list[Paragraph], list[str]]:
    """Return paragraphs for two files and their original names."""

    paragraphs: list[list[Paragraph]] = []
    filenames: list[str] = []
    for name, content in files:
        try:
            paragraphs.append(parse_docx_bytes(content))
            filenames.append(name)
        except Exception as exc:  # pragma: no cover - protective
            raise DocumentParseError(f"Unable to parse {name}") from exc
    if len(paragraphs) != 2:
        raise ValueError("Exactly two files are required")
    return paragraphs[0], paragraphs[1], filenames


@transaction.atomic
def store_diff_result(diff: DiffComputationResult, options: DiffOptions, filenames: Sequence[str]) -> DiffResult:
    """Persist a diff computation and return the saved instance."""

    result = DiffResult.objects.create(
        options=options.as_dict(),
        summary=diff.summary,
        diff_html=diff.html,
        diff_json=diff.to_json(),
        source_filenames=list(filenames),
    )
    return result


def perform_comparison(
    files: Iterable[tuple[str, bytes]], options: DiffOptions
) -> tuple[DiffComputationResult, list[str]]:
    """Parse two files, run the diff, and return the computation result."""

    paragraphs_a, paragraphs_b, filenames = parse_multiple(files)
    engine_options = EngineOptions(
        ignore_case=options.ignore_case,
        ignore_punctuation=options.ignore_punctuation,
        ignore_whitespace=options.ignore_whitespace,
    )
    url_options = options.url_options()
    return build_diff(paragraphs_a, paragraphs_b, engine_options, url_options), filenames


def render_diff_preview(diff: DiffComputationResult) -> str:
    """Return safe HTML for display in templates."""

    return diff.html


def export_diff_pdf(diff: DiffResult) -> bytes:
    """Generate a PDF document from stored HTML diff."""

    try:
        from weasyprint import HTML
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("WeasyPrint is not installed.") from exc

    body = f"""
    <html>
        <head>
            <meta charset='utf-8'>
            <style>
                body {{ font-family: sans-serif; color: #0f172a; }}
                .diff-view {{ display: flex; flex-direction: column; gap: 24px; }}
                .diff-view__legend {{ display: grid; grid-template-columns: repeat(2, 1fr); gap: 16px; font-size: 12px; font-weight: 600; text-transform: uppercase; color: #475569; }}
                .diff-panel {{ border: 1px solid #e2e8f0; border-radius: 16px; background: #ffffff; overflow: hidden; }}
                .diff-panel__header {{ display: flex; align-items: center; justify-content: space-between; padding: 12px 20px; border-bottom: 1px solid #e2e8f0; }}
                .diff-panel__meta {{ display: flex; align-items: center; gap: 12px; }}
                .diff-panel__icon {{ display: inline-flex; align-items: center; justify-content: center; width: 32px; height: 32px; border-radius: 999px; font-size: 16px; background: rgba(15, 23, 42, 0.08); }}
                .diff-panel__text {{ display: flex; flex-direction: column; gap: 4px; font-size: 14px; }}
                .diff-panel__label {{ font-weight: 600; }}
                .diff-panel__description {{ font-size: 12px; color: #475569; }}
                .diff-panel__anchor {{ color: #94a3b8; font-size: 12px; text-decoration: none; }}
                .diff-panel__columns {{ display: grid; grid-template-columns: repeat(2, 1fr); }}
                .diff-panel__column {{ padding: 20px; font-size: 14px; line-height: 1.6; border-top: 1px solid #e2e8f0; background: #ffffff; }}
                .diff-panel__column:first-child {{ border-right: 1px solid #e2e8f0; border-top: 0; }}
                .diff-panel__column-title {{ display: block; margin-bottom: 8px; font-size: 12px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.08em; color: #94a3b8; }}
                .diff-panel--equal .diff-panel__header {{ background: #f8fafc; color: #475569; }}
                .diff-panel--equal .diff-panel__column {{ background: #f8fafc; }}
                .diff-panel--equal .diff-panel__icon {{ background: #e2e8f0; color: #475569; }}
                .diff-panel--insert .diff-panel__header {{ background: #ecfdf5; color: #047857; }}
                .diff-panel--insert .diff-panel__column--right {{ background: #f0fdf4; }}
                .diff-panel--insert .diff-panel__icon {{ background: #bbf7d0; color: #047857; }}
                .diff-panel--delete .diff-panel__header {{ background: #fef2f2; color: #b91c1c; }}
                .diff-panel--delete .diff-panel__column--left {{ background: #fef2f2; }}
                .diff-panel--delete .diff-panel__icon {{ background: #fecaca; color: #b91c1c; }}
                .diff-panel--replace .diff-panel__header {{ background: #fffbeb; color: #b45309; }}
                .diff-panel--replace .diff-panel__column {{ background: #fffbeb; }}
                .diff-panel--replace .diff-panel__icon {{ background: #fde68a; color: #b45309; }}
                .diff-ins {{ background-color: #d1fae5; }}
                .diff-del {{ background-color: #fee2e2; }}
            </style>
        </head>
        <body>
            <h1>Document Diff</h1>
            <pre>{escape(json.dumps(diff.summary, indent=2))}</pre>
            {diff.diff_html}
        </body>
    </html>
    """
    return HTML(string=body).write_pdf()
